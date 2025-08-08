import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
import traceback

class ResumeGenerator:
    def __init__(self):
        self.templates = {
            'modern': self._generate_modern_resume,
            'classic': self._generate_classic_resume,
            'minimal': self._generate_minimal_resume
        }
        
        # Create generated_resumes directory if it doesn't exist
        self.output_dir = 'generated_resumes'
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
    
    def generate_resume(self, profile_data, template='modern', format='docx'):
        """Generate resume in specified format"""
        try:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            
            if format == 'docx':
                filename = f'resume_{template}_{timestamp}.docx'
                filepath = os.path.join(self.output_dir, filename)
                
                print(f"[DEBUG] Generating DOCX resume: {filepath}")
                print(f"[DEBUG] Profile data keys: {list(profile_data.keys())}")
                
                if template in self.templates:
                    self.templates[template](profile_data, filepath)
                else:
                    print(f"[ERROR] Unknown template: {template}")
                    raise ValueError(f"Unknown template: {template}")
                    
            elif format == 'pdf':
                filename = f'resume_{template}_{timestamp}.pdf'
                filepath = os.path.join(self.output_dir, filename)
                
                print(f"[DEBUG] Generating PDF resume: {filepath}")
                
                if template in self.templates:
                    self._generate_pdf_resume(profile_data, filepath, template)
                else:
                    print(f"[ERROR] Unknown template: {template}")
                    raise ValueError(f"Unknown template: {template}")
            else:
                raise ValueError(f"Unsupported format: {format}")
            
            # Verify file was created
            if os.path.exists(filepath):
                print(f"[DEBUG] File created successfully: {filepath}")
                print(f"[DEBUG] File size: {os.path.getsize(filepath)} bytes")
                return filename
            else:
                print(f"[ERROR] File was not created: {filepath}")
                raise FileNotFoundError(f"Resume file was not created: {filepath}")
                
        except Exception as e:
            print(f"[ERROR] Failed to generate resume: {e}")
            print(f"[ERROR] Traceback: {traceback.format_exc()}")
            raise
    
    def _get_field(self, profile_data, key):
        if 'personal_info' in profile_data and key in profile_data['personal_info']:
            return profile_data['personal_info'][key]
        return profile_data.get(key, '')

    def _generate_modern_resume(self, profile_data, filepath):
        """Generate modern style resume"""
        doc = Document()
        
        # Set page margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Header with name and contact info
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Name
        name = self._get_field(profile_data, 'name')
        name_run = header.add_run(name)
        name_run.font.size = Pt(24)
        name_run.font.bold = True
        name_run.font.name = 'Calibri'
        
        # Headline
        doc.add_paragraph()
        headline = self._get_field(profile_data, 'headline')
        headline_para = doc.add_paragraph(headline)
        headline_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if headline_para.runs:
            headline_para.runs[0].font.size = Pt(14)
            headline_para.runs[0].font.italic = True
            headline_para.runs[0].font.name = 'Calibri'
        
        # Contact information
        email = self._get_field(profile_data, 'email')
        phone = self._get_field(profile_data, 'phone')
        location = self._get_field(profile_data, 'location')
        contact_info = doc.add_paragraph()
        contact_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_info.add_run(f"{email} | {phone} | {location}")
        if contact_info.runs:
            contact_info.runs[0].font.size = Pt(10)
            contact_info.runs[0].font.name = 'Calibri'
        
        doc.add_paragraph()  # Spacing
        
        # Summary
        summary = self._get_field(profile_data, 'summary')
        if summary:
            self._add_section_header(doc, "PROFESSIONAL SUMMARY")
            summary_para = doc.add_paragraph(summary)
            if summary_para.runs:
                summary_para.runs[0].font.size = Pt(11)
                summary_para.runs[0].font.name = 'Calibri'
            doc.add_paragraph()  # Spacing
        
        # Experience
        experience = profile_data.get('experience', [])
        if experience:
            self._add_section_header(doc, "PROFESSIONAL EXPERIENCE")
            for exp in experience:
                # Job title and company
                job_header = doc.add_paragraph()
                job_title = job_header.add_run(f"{exp.get('title', '')} - {exp.get('company', '')}")
                job_title.font.bold = True
                job_title.font.size = Pt(12)
                job_title.font.name = 'Calibri'
                
                # Duration and location
                duration_location = doc.add_paragraph()
                duration_location.add_run(f"{exp.get('duration', '')} | {exp.get('location', '')}")
                if duration_location.runs:
                    duration_location.runs[0].font.italic = True
                    duration_location.runs[0].font.size = Pt(10)
                    duration_location.runs[0].font.name = 'Calibri'
                
                # Description
                description = exp.get('description', '')
                if description:
                    description_para = doc.add_paragraph(description)
                    if description_para.runs:
                        description_para.runs[0].font.size = Pt(11)
                        description_para.runs[0].font.name = 'Calibri'
                doc.add_paragraph()  # Spacing
        
        # Education
        education = profile_data.get('education', [])
        if education:
            self._add_section_header(doc, "EDUCATION")
            for edu in education:
                # Degree and school
                edu_header = doc.add_paragraph()
                degree_school = edu_header.add_run(f"{edu.get('degree', '')} - {edu.get('school', '')}")
                degree_school.font.bold = True
                degree_school.font.size = Pt(12)
                degree_school.font.name = 'Calibri'
                
                # Duration and GPA
                edu_details = doc.add_paragraph()
                details_text = f"{edu.get('duration', '')} | {edu.get('location', '')}"
                if edu.get('gpa'):
                    details_text += f" | GPA: {edu['gpa']}"
                edu_details.add_run(details_text)
                if edu_details.runs:
                    edu_details.runs[0].font.italic = True
                    edu_details.runs[0].font.size = Pt(10)
                    edu_details.runs[0].font.name = 'Calibri'
                doc.add_paragraph()  # Spacing
        
        # Skills
        skills = profile_data.get('skills', [])
        if skills:
            self._add_section_header(doc, "TECHNICAL SKILLS")
            skills_text = ", ".join(skills)
            skills_para = doc.add_paragraph(skills_text)
            if skills_para.runs:
                skills_para.runs[0].font.size = Pt(11)
                skills_para.runs[0].font.name = 'Calibri'
            doc.add_paragraph()  # Spacing
        
        # Save document
        print(f"[DEBUG] Saving resume to: {filepath}")
        try:
            doc.save(filepath)
            print(f"[DEBUG] Resume saved successfully: {filepath}")
        except Exception as e:
            print(f"[ERROR] Failed to save resume: {e}")
    
    def _add_section_header(self, doc, title):
        """Add section header for modern template"""
        header = doc.add_paragraph()
        header_run = header.add_run(title)
        header_run.font.bold = True
        header_run.font.size = Pt(14)
        header_run.font.name = 'Calibri'
        header_run.font.underline = True
        
        doc.add_paragraph()  # Spacing 

    def _generate_classic_resume(self, profile_data, filepath):
        # Placeholder: use modern resume for now
        self._generate_modern_resume(profile_data, filepath)

    def _generate_minimal_resume(self, profile_data, filepath):
        # Placeholder: use modern resume for now
        self._generate_modern_resume(profile_data, filepath) 

    def _generate_pdf_resume(self, profile_data, filepath, template):
        """Generate PDF resume"""
        try:
            doc = SimpleDocTemplate(filepath, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # Create custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                spaceAfter=12,
                alignment=TA_CENTER,
                textColor=colors.darkblue
            )
            
            subtitle_style = ParagraphStyle(
                'CustomSubtitle',
                parent=styles['Normal'],
                fontSize=14,
                spaceAfter=6,
                alignment=TA_CENTER,
                textColor=colors.grey
            )
            
            section_style = ParagraphStyle(
                'SectionHeader',
                parent=styles['Heading2'],
                fontSize=16,
                spaceAfter=6,
                textColor=colors.darkblue
            )
            
            # Header
            name = profile_data.get('personal_info', {}).get('name') or profile_data.get('name', '')
            headline = profile_data.get('personal_info', {}).get('headline') or profile_data.get('headline', '')
            email = profile_data.get('personal_info', {}).get('email') or profile_data.get('email', '')
            phone = profile_data.get('personal_info', {}).get('phone') or profile_data.get('phone', '')
            location = profile_data.get('personal_info', {}).get('location') or profile_data.get('location', '')
            
            story.append(Paragraph(name, title_style))
            if headline:
                story.append(Paragraph(headline, subtitle_style))
            
            # Contact info
            contact_info = f"{email} | {phone} | {location}"
            story.append(Paragraph(contact_info, subtitle_style))
            story.append(Spacer(1, 12))
            
            # Summary
            summary = profile_data.get('personal_info', {}).get('summary') or profile_data.get('summary', '')
            if summary:
                story.append(Paragraph("PROFESSIONAL SUMMARY", section_style))
                story.append(Paragraph(summary, styles['Normal']))
                story.append(Spacer(1, 12))
            
            # Work Experience
            experience = profile_data.get('experience', [])
            if experience:
                story.append(Paragraph("WORK EXPERIENCE", section_style))
                for exp in experience:
                    title = exp.get('title', '')
                    company = exp.get('company', '')
                    duration = exp.get('duration', '')
                    description = exp.get('description', '')
                    
                    exp_header = f"<b>{title}</b> - {company} ({duration})"
                    story.append(Paragraph(exp_header, styles['Normal']))
                    if description:
                        story.append(Paragraph(description, styles['Normal']))
                    story.append(Spacer(1, 6))
                story.append(Spacer(1, 12))
            
            # Education
            education = profile_data.get('education', [])
            if education:
                story.append(Paragraph("EDUCATION", section_style))
                for edu in education:
                    degree = edu.get('degree', '')
                    institution = edu.get('institution', '')
                    year = edu.get('year', '')
                    
                    edu_text = f"<b>{degree}</b> - {institution} ({year})"
                    story.append(Paragraph(edu_text, styles['Normal']))
                    story.append(Spacer(1, 6))
                story.append(Spacer(1, 12))
            
            # Skills
            skills = profile_data.get('skills', [])
            if skills:
                story.append(Paragraph("SKILLS", section_style))
                skills_text = ", ".join(skills)
                story.append(Paragraph(skills_text, styles['Normal']))
                story.append(Spacer(1, 12))
            
            # Build PDF
            doc.build(story)
            print(f"[DEBUG] PDF generated successfully: {filepath}")
            
        except Exception as e:
            print(f"[ERROR] Failed to generate PDF: {e}")
            print(f"[ERROR] PDF Traceback: {traceback.format_exc()}")
            raise 